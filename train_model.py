import numpy as np
import pickle as pkl
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from nltk.stem import WordNetLemmatizer
import re
import docx
from typing import Dict, List, Tuple, Set
import logging
from pathlib import Path
from collections import defaultdict

nltk.download('punkt', quiet=True)
nltk.download('punkt_tab', quiet=True)
nltk.download('stopwords', quiet=True)
nltk.download('wordnet', quiet=True)

logging.basicConfig(level=logging.INFO)

class DocumentComplianceAnalyzer:
    compliance_threshold = 0.7

    def __init__(self):
        self.vectorizer = TfidfVectorizer(
            min_df=1,
            ngram_range=(1, 2),
            stop_words='english',
            lowercase=True,
            strip_accents='unicode',
            norm='l2'
        )
        self.lemmatizer = WordNetLemmatizer()
        self.stop_words = set(stopwords.words('english'))

    def extract_text_from_docx(self, docx_path: str) -> str:
        """Extract and concatenate all text from a Word document."""
        if not Path(docx_path).is_file():
            logging.error(f"File not found: {docx_path}")
            return ""
        try:
            doc = docx.Document(docx_path)
            return "\n".join(para.text for para in doc.paragraphs if para.text.strip())
        except Exception as e:
            logging.error(f"Error extracting text from {docx_path}: {e}")
            return ""

    def extract_headings_from_docx(self, docx_path: str) -> List[Tuple[str, int]]:
        """Extract headings with their level (Heading 1, Heading 2, etc.)"""
        try:
            doc = docx.Document(docx_path)
            headings = []
            for para in doc.paragraphs:
                if para.style.name.startswith("Heading"):
                    try:
                        level = int(para.style.name.split(" ")[1])
                    except (IndexError, ValueError):
                        level = 1  # Default to level 1 if can't parse
                    headings.append((para.text.strip(), level))
            return headings
        except Exception as e:
            logging.error(f"Error extracting headings from {docx_path}: {e}")
            return []

    def extract_formatting_from_docx(self, docx_path: str) -> Dict:
        """Extract comprehensive formatting information with frequencies."""
        try:
            doc = docx.Document(docx_path)
            formatting = {
                "fonts": defaultdict(int),
                "sizes": defaultdict(int),
                "bold_count": 0,
                "italic_count": 0,
                "underline_count": 0,
                "alignments": defaultdict(int),
                "styles": defaultdict(int),
                "line_spacing": defaultdict(int),
                "paragraph_spacing": defaultdict(int)
            }
            
            for para in doc.paragraphs:
                formatting["styles"][para.style.name] += 1
                if para.paragraph_format:
                    formatting["line_spacing"][round(para.paragraph_format.line_spacing or 1.0, 1)] += 1
                    spacing = round((para.paragraph_format.space_before or 0) + 
                                  (para.paragraph_format.space_after or 0), 1)
                    formatting["paragraph_spacing"][spacing] += 1
                
                for run in para.runs:
                    if run.font:
                        font_name = run.font.name or "Unknown"
                        formatting["fonts"][font_name] += len(run.text)
                        size = run.font.size.pt if run.font.size else 0
                        formatting["sizes"][size] += len(run.text)
                        formatting["bold_count"] += len(run.text) if run.bold else 0
                        formatting["italic_count"] += len(run.text) if run.italic else 0
                        formatting["underline_count"] += len(run.text) if run.underline else 0
                
                if para.alignment is not None:
                    formatting["alignments"][para.alignment] += 1
            
            # Convert defaultdict to regular dict for JSON serialization
            for key in ["fonts", "sizes", "alignments", "styles", "line_spacing", "paragraph_spacing"]:
                formatting[key] = dict(formatting[key])
                
            return formatting
        except Exception as e:
            logging.error(f"Error extracting formatting from {docx_path}: {e}")
            return {}

    def preprocess_text(self, text: str) -> str:
        """Clean and normalize text while preserving structure."""
        text = text.lower().strip()
        text = re.sub(r'[^\w\s.,!?]', '', text)  # Keep basic punctuation
        text = re.sub(r'\s+', ' ', text)  # Normalize whitespace
        tokens = word_tokenize(text)
        processed = [self.lemmatizer.lemmatize(token) for token in tokens if token not in self.stop_words]
        return ' '.join(processed)

    def calculate_similarity(self, doc1: str, doc2: str) -> float:
        """Calculate text similarity with input validation."""
        if not hasattr(self.vectorizer, "vocabulary_"):
            logging.error("TF-IDF vectorizer is not fitted.")
            return 0.0
        
        if not doc1 or not doc2:
            return 0.0
            
        doc1_processed = self.preprocess_text(doc1)
        doc2_processed = self.preprocess_text(doc2)
        
        try:
            tfidf_matrix = self.vectorizer.transform([doc1_processed, doc2_processed])
            return float(cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0])
        except Exception as e:
            logging.error(f"Error calculating similarity: {e}")
            return 0.0

    def _compare_headings(self, student: List[Tuple[str, int]], template: List[Tuple[str, int]]) -> float:
        """Compare heading structure with level awareness."""
        if not template:
            return 1.0  # No headings to compare against
            
        # Compare both text and hierarchy
        matches = 0
        template_dict = {text.lower(): level for text, level in template}
        
        for s_text, s_level in student:
            s_text_lower = s_text.lower()
            if s_text_lower in template_dict:
                if s_level == template_dict[s_text_lower]:
                    matches += 1  # Perfect match
                else:
                    matches += 0.5  # Right text but wrong level
        
        return matches / len(template)

    def _compare_formatting(self, student: Dict, template: Dict) -> float:
        """Comprehensive formatting comparison with weighted scores."""
        scores = []
        
        # Font comparison (weighted by usage frequency)
        font_score = self._compare_frequency_dist(student["fonts"], template["fonts"])
        scores.append(("fonts", 0.2, font_score))
        
        # Size comparison
        size_score = self._compare_frequency_dist(student["sizes"], template["sizes"])
        scores.append(("sizes", 0.2, size_score))
        
        # Style attributes comparison
        style_score = (
            0.7 * min(1, student["bold_count"] / max(1, template["bold_count"])) +
            0.2 * min(1, student["italic_count"] / max(1, template["italic_count"])) +
            0.1 * min(1, student["underline_count"] / max(1, template["underline_count"]))
        )
        scores.append(("styles", 0.2, style_score))
        
        # Alignment comparison
        align_score = self._compare_frequency_dist(student["alignments"], template["alignments"])
        scores.append(("alignment", 0.1, align_score))
        
        # Paragraph styles
        para_score = (
            0.5 * self._compare_frequency_dist(student["line_spacing"], template["line_spacing"]) +
            0.5 * self._compare_frequency_dist(student["paragraph_spacing"], template["paragraph_spacing"])
        )
        scores.append(("paragraph", 0.2, para_score))
        
        # Overall style distribution
        style_dist_score = self._compare_frequency_dist(student["styles"], template["styles"])
        scores.append(("style_dist", 0.1, style_dist_score))
        
        # Calculate weighted average
        total_weight = sum(weight for _, weight, _ in scores)
        weighted_sum = sum(weight * score for _, weight, score in scores)
        
        return weighted_sum / total_weight if total_weight > 0 else 0

    def _compare_frequency_dist(self, student: Dict, template: Dict) -> float:
        """Compare frequency distributions of formatting attributes."""
        if not template:
            return 1.0
            
        total = sum(template.values())
        match = 0
        
        for attr, count in student.items():
            if attr in template:
                match += min(count, template[attr])
        
        return match / total

    def evaluate_compliance(self, student_doc: str, template_doc: str, 
                          student_format: Dict, template_format: Dict, 
                          student_headings: List[Tuple[str, int]], 
                          template_headings: List[Tuple[str, int]]) -> Dict:
        """Comprehensive compliance evaluation with detailed feedback."""
        if not student_doc or not template_doc:
            logging.error("Empty document(s) provided")
            return {
                "text_similarity": 0,
                "heading_compliance": 0,
                "formatting_compliance": 0,
                "final_compliance_score": 0,
                "is_compliant": False,
                "recommendations": ["Empty document provided"]
            }

        # Calculate similarity with smoothing
        raw_similarity = self.calculate_similarity(student_doc, template_doc)
        similarity_score = 0.3 + 0.7 * raw_similarity  # Reduce harsh penalties
        
        # Compare headings with hierarchy awareness
        heading_score = self._compare_headings(student_headings, template_headings)
        
        # Comprehensive formatting comparison
        formatting_score = self._compare_formatting(student_format, template_format)
        
        # Calculate final score with balanced weights
        final_score = (
            similarity_score * 0.4 +
            heading_score * 0.3 +
            formatting_score * 0.3
        )
        
        # Generate detailed recommendations
        recommendations = self._generate_recommendations(
            similarity_score, 
            heading_score, 
            formatting_score,
            student_format,
            template_format,
            student_headings,
            template_headings
        )
        
        return {
            "text_similarity": raw_similarity,
            "heading_compliance": heading_score,
            "formatting_compliance": formatting_score,
            "final_compliance_score": final_score,
            "is_compliant": final_score >= self.compliance_threshold,
            "recommendations": recommendations
        }

    def _generate_recommendations(self, similarity: float, heading: float, 
                                formatting: float, student_fmt: Dict, 
                                template_fmt: Dict, student_hdgs: List[Tuple[str, int]], 
                                template_hdgs: List[Tuple[str, int]]) -> List[str]:
        """Generate specific, actionable recommendations."""
        recommendations = []
        
        # Content recommendations
        if similarity < 0.5:
            recommendations.append("Significant content differences detected - review template")
        elif similarity < 0.7:
            recommendations.append("Moderate content differences - consider closer alignment")
        elif similarity < 0.9:
            recommendations.append("Minor content differences detected")
        
        # Heading recommendations
        missing_headings = [h[0] for h in template_hdgs 
                          if h[0].lower() not in {sh[0].lower() for sh in student_hdgs}]
        if missing_headings:
            recommendations.append(f"Missing headings: {', '.join(missing_headings[:3])}{'...' if len(missing_headings)>3 else ''}")
            
        level_mismatches = [h[0] for h in template_hdgs 
                          for sh in student_hdgs 
                          if h[0].lower() == sh[0].lower() and h[1] != sh[1]]
        if level_mismatches:
            recommendations.append(f"Heading level mismatches: {', '.join(level_mismatches[:3])}")
        
        # Formatting recommendations
        fmt_issues = []
        
        # Font issues
        main_template_font = max(template_fmt["fonts"].items(), key=lambda x: x[1], default=None)
        main_student_font = max(student_fmt["fonts"].items(), key=lambda x: x[1], default=None)
        if main_template_font and main_student_font and main_template_font[0] != main_student_font[0]:
            fmt_issues.append(f"main font ({main_student_font[0]} vs template {main_template_font[0]})")
        
        # Size issues
        main_template_size = max(template_fmt["sizes"].items(), key=lambda x: x[1], default=None)
        main_student_size = max(student_fmt["sizes"].items(), key=lambda x: x[1], default=None)
        if main_template_size and main_student_size and main_template_size[0] != main_student_size[0]:
            fmt_issues.append(f"font size ({main_student_size[0]}pt vs template {main_template_size[0]}pt)")
        
        # Style issues
        style_diffs = []
        if student_fmt["bold_count"] < 0.8 * template_fmt["bold_count"]:
            style_diffs.append("not enough bold text")
        elif student_fmt["bold_count"] > 1.2 * template_fmt["bold_count"]:
            style_diffs.append("too much bold text")
            
        if student_fmt["italic_count"] < 0.8 * template_fmt["italic_count"]:
            style_diffs.append("not enough italic text")
        elif student_fmt["italic_count"] > 1.2 * template_fmt["italic_count"]:
            style_diffs.append("too much italic text")
            
        if style_diffs:
            fmt_issues.append(f"style issues: {', '.join(style_diffs)}")
        
        if fmt_issues:
            recommendations.append("Formatting differences: " + ", ".join(fmt_issues))
        
        return recommendations

    def save_model(self, path: str):
        """Save the trained vectorizer to disk."""
        with open(path, 'wb') as f:
            pkl.dump({
                'vectorizer': self.vectorizer,
                'compliance_threshold': self.compliance_threshold
            }, f)

    def load_model(self, path: str):
        """Load a trained vectorizer from disk."""
        with open(path, 'rb') as f:
            data = pkl.load(f)
            self.vectorizer = data['vectorizer']
            self.compliance_threshold = data.get('compliance_threshold', 0.7)

    def fit_vectorizer(self, corpus: List[str]):
        """Fit the vectorizer to a corpus of documents."""
        self.vectorizer.fit(corpus)

if __name__ == "__main__":
    analyzer = DocumentComplianceAnalyzer()

    template_path = Path("C:/Users/USER/Documents/Test_docs/Letter-Document-Template.docx")
    student_path = Path("C:/Users/USER/Documents/Test_docs/Letter-Document-Template (1).docx")

    # Extract content
    template_text = analyzer.extract_text_from_docx(template_path)
    student_text = analyzer.extract_text_from_docx(student_path)

    # Train and save model
    analyzer.fit_vectorizer([template_text, student_text])
    analyzer.save_model("vectorizer.pkl")

    # Load and analyze
    analyzer.load_model("vectorizer.pkl")

    template_format = analyzer.extract_formatting_from_docx(template_path)
    student_format = analyzer.extract_formatting_from_docx(student_path)
    template_headings = analyzer.extract_headings_from_docx(template_path)
    student_headings = analyzer.extract_headings_from_docx(student_path)

    compliance_result = analyzer.evaluate_compliance(
        student_text, template_text,
        student_format, template_format,
        student_headings, template_headings
    )

    print("\nCompliance Analysis Results:")
    print(f"Text Similarity: {compliance_result['text_similarity']:.2f}")
    print(f"Heading Compliance: {compliance_result['heading_compliance']:.2f}")
    print(f"Formatting Compliance: {compliance_result['formatting_compliance']:.2f}")
    print(f"Final Compliance Score: {compliance_result['final_compliance_score']:.2f}")
    print(f"Document is {'COMPLIANT' if compliance_result['is_compliant'] else 'NON-COMPLIANT'}")
    
    if compliance_result['recommendations']:
        print("\nRecommendations for Improvement:")
        for i, rec in enumerate(compliance_result['recommendations'], 1):
            print(f"{i}. {rec}")